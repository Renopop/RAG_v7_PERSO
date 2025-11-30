# docx_processing.py
import docx
import re
import os
import tempfile
import shutil
import logging


def convert_doc_to_docx(doc_path: str) -> str:
    """
    Convertit un fichier .doc en .docx via Microsoft Word (Windows).

    Requiert:
    - Windows avec Microsoft Word installé
    - pywin32 (pip install pywin32)

    Traitement automatique:
    - Accepte toutes les révisions (suivi des modifications)
    - Supprime tous les commentaires
    - Stocke le .docx dans un dossier temporaire (l'original n'est pas modifié)

    Retourne le chemin vers le fichier .docx créé (dans un dossier temporaire).
    Le fichier temporaire doit être supprimé par l'appelant après usage.
    """
    if not os.path.isfile(doc_path):
        raise FileNotFoundError(f"Fichier introuvable: {doc_path}")

    if os.name != 'nt':
        raise RuntimeError(
            "La conversion .doc → .docx nécessite Windows avec Microsoft Word. "
            "Sur Linux/Mac, convertir manuellement les fichiers .doc en .docx."
        )

    try:
        import win32com.client
        import pythoncom
    except ImportError:
        raise RuntimeError(
            "pywin32 non installé. Exécutez: pip install pywin32\n"
            "Puis redémarrez Python/l'application."
        )

    temp_dir = tempfile.mkdtemp(prefix="doc_convert_")
    logging.info(f"[docx_processing] Conversion .doc → .docx: {os.path.basename(doc_path)}")

    # Initialiser COM pour ce thread
    pythoncom.CoInitialize()

    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False

        # Chemin absolu requis par Word
        abs_doc_path = os.path.abspath(doc_path)

        # Ouvrir le document en lecture seule pour ne pas modifier l'original
        # ReadOnly=True empêche la modification du fichier source
        doc = word.Documents.Open(abs_doc_path, ReadOnly=True)

        # Accepter toutes les révisions (suivi des modifications)
        if doc.Revisions.Count > 0:
            logging.info(f"[docx_processing] Acceptation de {doc.Revisions.Count} révision(s)")
            doc.AcceptAllRevisions()

        # Supprimer tous les commentaires
        if doc.Comments.Count > 0:
            logging.info(f"[docx_processing] Suppression de {doc.Comments.Count} commentaire(s)")
            # Supprimer du dernier au premier pour éviter les problèmes d'index
            for i in range(doc.Comments.Count, 0, -1):
                doc.Comments(i).Delete()

        # Chemin de sortie (fichier temporaire)
        base_name = os.path.splitext(os.path.basename(doc_path))[0]
        docx_path = os.path.join(temp_dir, f"{base_name}.docx")
        abs_docx_path = os.path.abspath(docx_path)

        # Sauvegarder en .docx (format 16 = wdFormatXMLDocument)
        # Le fichier est sauvegardé dans le dossier temporaire, pas sur l'original
        doc.SaveAs2(abs_docx_path, FileFormat=16)

        logging.info(f"[docx_processing] Conversion Word réussie: {docx_path}")
        return docx_path

    except Exception as e:
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise RuntimeError(
            f"Erreur lors de la conversion avec Word: {e}\n"
            f"Vérifiez que Microsoft Word est installé et fonctionne."
        )

    finally:
        if doc:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def docx_to_text(docx_path):
    """
    Extrait le texte brut d'un fichier .docx en conservant les sauts de ligne entre paragraphes.
    """
    doc = docx.Document(docx_path)
    full_text = [para.text for para in doc.paragraphs]
    # On garde un \n entre chaque paragraphe
    return "\n".join(full_text)

def normalize_whitespace(text: str) -> str:
    """
    Normalise les espaces dans le texte EN CONSERVANT les sauts de ligne.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")
    norm_lines = []
    for line in lines:
        # Réduire les suites d'espaces/tabulations
        line = re.sub(r"[ \t]+", " ", line)
        norm_lines.append(line.strip())
    return "\n".join(norm_lines).strip()

def extract_text_from_docx(file_path: str) -> str:
    """
    Extrait et nettoie le texte d'un fichier DOCX ou DOC.

    Pour les fichiers .doc (ancien format), convertit automatiquement en .docx
    via Microsoft Word avant extraction (Windows uniquement).

    Args:
        file_path: Chemin vers le fichier .doc ou .docx

    Returns:
        Texte extrait et nettoyé
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".doc":
        # Convertir .doc → .docx avec Microsoft Word
        temp_docx = None
        try:
            temp_docx = convert_doc_to_docx(file_path)
            raw_text = docx_to_text(temp_docx)
        finally:
            # Nettoyer le fichier temporaire
            if temp_docx and os.path.isfile(temp_docx):
                temp_dir = os.path.dirname(temp_docx)
                shutil.rmtree(temp_dir, ignore_errors=True)
    else:
        # Fichier .docx natif
        raw_text = docx_to_text(file_path)

    return normalize_whitespace(raw_text)

def extract_paragraphs_from_docx(docx_path):
    """
    Extrait les paragraphes individuels d'un fichier DOCX.
    """
    doc = docx.Document(docx_path)
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return paragraphs

def extract_sections_from_docx(docx_path, heading_styles=None):
    """
    Extrait les sections d'un fichier DOCX en se basant sur les styles de titre.
    heading_styles est une liste de noms de styles considérés comme des titres de sections.
    """
    if heading_styles is None:
        heading_styles = ["Heading 1", "Heading 2", "Titre 1", "Titre 2"]

    doc = docx.Document(docx_path)
    sections = []
    current_section_title = None
    current_section_content = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""

        if style_name in heading_styles:
            # Sauvegarder la section précédente
            if current_section_title or current_section_content:
                sections.append({
                    "title": current_section_title,
                    "content": "\n".join(current_section_content).strip(),
                })
                current_section_content = []

            # Nouveau titre de section
            current_section_title = para.text.strip()
        else:
            # Contenu de la section en cours
            if para.text.strip():
                current_section_content.append(para.text.strip())

    # Ajouter la dernière section si présente
    if current_section_title or current_section_content:
        sections.append({
            "title": current_section_title,
            "content": "\n".join(current_section_content).strip(),
        })

    return sections

def extract_text_from_tables(docx_path):
    """
    Extrait le texte contenu dans les tableaux d'un fichier DOCX.
    """
    doc = docx.Document(docx_path)
    table_texts = []

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                table_texts.append(" | ".join(row_text))

    return "\n".join(table_texts)
