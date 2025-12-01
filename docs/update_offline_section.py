#!/usr/bin/env python3
"""
Script pour ajouter la section Mode OFFLINE au GUIDE_RAG.docx
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level):
    """Ajoute un titre."""
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False):
    """Ajoute un paragraphe."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet_list(doc, items):
    """Ajoute une liste a puces."""
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def add_table(doc, headers, rows):
    """Ajoute un tableau."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text

    return table

def add_code_block(doc, code_text):
    """Ajoute un bloc de code."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def main():
    # Ouvrir le document existant
    doc = Document('GUIDE_RAG.docx')

    # Ajouter une page de saut avant la nouvelle section
    doc.add_page_break()

    # ==========================================================================
    # SECTION: MODE OFFLINE
    # ==========================================================================

    add_heading(doc, "15. Mode OFFLINE (v2.1)", level=1)

    add_paragraph(doc, """Le mode offline permet d'utiliser le systeme RAG sans connexion internet,
en utilisant des modeles IA locaux sur GPU NVIDIA. Cette fonctionnalite est ideale pour
les environnements deconnectes ou les situations ou la latence reseau est critique.""")

    doc.add_paragraph()

    # 15.1 Prerequisites
    add_heading(doc, "15.1 Prerequisites materiel", level=2)

    add_table(doc,
        ["Composant", "Minimum", "Recommande"],
        [
            ["GPU", "RTX 3080 (10 GB)", "RTX 4090 (24 GB)"],
            ["VRAM", "12 GB", "16+ GB"],
            ["Disque", "20 GB (modeles)", "SSD NVMe"],
        ]
    )

    doc.add_paragraph()

    # 15.2 Modeles locaux
    add_heading(doc, "15.2 Modeles locaux utilises", level=2)

    add_paragraph(doc, "Le mode offline utilise trois modeles IA locaux :")
    doc.add_paragraph()

    add_table(doc,
        ["Modele", "Fonction", "VRAM utilisee"],
        [
            ["BGE-M3", "Embeddings (1024 dimensions)", "~2 GB"],
            ["BGE-Reranker-v2-m3", "Re-ranking des resultats", "~2 GB"],
            ["Mistral-7B-Instruct-v0.3", "Generation de reponses", "~8 GB"],
        ]
    )

    doc.add_paragraph()

    # 15.3 Activation
    add_heading(doc, "15.3 Activation du mode offline", level=2)

    add_paragraph(doc, "Pour activer le mode offline :")
    add_bullet_list(doc, [
        "1. Dans la sidebar de l'application, section 'Mode de fonctionnement'",
        "2. Cochez 'Mode OFFLINE (modeles locaux)'",
        "3. Un spinner 'Chargement des modeles IA locaux...' apparait",
        "4. Les modeles sont pre-charges en memoire GPU",
    ])

    doc.add_paragraph()

    # 15.4 Pre-chargement
    add_heading(doc, "15.4 Pre-chargement automatique", level=2)

    add_paragraph(doc, """Au demarrage en mode offline, les modeles sont pre-charges automatiquement
pour eviter les temps de chargement pendant les requetes. Voici un exemple de sortie :""")

    doc.add_paragraph()

    preload_output = """============================================================
[PRELOAD] Pre-chargement des modeles offline...
============================================================
[PRELOAD] 1/3 Chargement BGE-M3 (embeddings)...
[PRELOAD] BGE-M3 charge en 3.2s
[PRELOAD] 2/3 Chargement BGE-Reranker...
[PRELOAD] BGE-Reranker charge en 2.1s
[PRELOAD] 3/3 Chargement Mistral-7B (LLM)...
[PRELOAD] Mistral-7B charge en 24.5s
============================================================
[PRELOAD] Tous les modeles charges (29.8s)
[PRELOAD] VRAM utilisee: 12.5/24.0 GB
============================================================"""

    add_code_block(doc, preload_output)

    doc.add_paragraph()

    add_paragraph(doc, "Avantages du pre-chargement :", bold=True)
    add_bullet_list(doc, [
        "Les requetes sont plus rapides (pas de temps de chargement)",
        "Les modeles restent en memoire GPU",
        "Le chargement se fait une seule fois au demarrage",
    ])

    doc.add_paragraph()

    # 15.5 Performance
    add_heading(doc, "15.5 Performance en mode offline (RTX 4090)", level=2)

    add_table(doc,
        ["Operation", "Temps"],
        [
            ["Recherche Hybrid Search (5000 chunks)", "~8s"],
            ["BGE Reranker (30 documents)", "~3s"],
            ["Generation de reponse LLM", "~7s"],
            ["Total par requete", "~20s"],
        ]
    )

    doc.add_paragraph()

    # 15.6 Emplacement des modeles
    add_heading(doc, "15.6 Emplacement des modeles", level=2)

    add_paragraph(doc, "Les modeles doivent etre telecharges dans les chemins suivants :")
    doc.add_paragraph()

    add_table(doc,
        ["Modele", "Chemin"],
        [
            ["BGE-M3", "D:\\LLM_Models\\bge-m3"],
            ["BGE-Reranker", "D:\\LLM_Models\\bge-reranker-v2-m3"],
            ["Mistral-7B", "D:\\LLM_Models\\Mistral-7B-Instruct-v0.3"],
        ]
    )

    doc.add_paragraph()
    add_paragraph(doc, "Note : Les chemins sont configurables dans config.json ou via la page de configuration.")

    doc.add_paragraph()

    # 15.7 Telechargement
    add_heading(doc, "15.7 Telechargement des modeles", level=2)

    add_paragraph(doc, "Les modeles peuvent etre telecharges depuis Hugging Face avec les commandes suivantes :")
    doc.add_paragraph()

    download_commands = """# BGE-M3 (embeddings)
huggingface-cli download BAAI/bge-m3 --local-dir D:\\LLM_Models\\bge-m3

# BGE-Reranker
huggingface-cli download BAAI/bge-reranker-v2-m3 --local-dir D:\\LLM_Models\\bge-reranker-v2-m3

# Mistral-7B-Instruct
huggingface-cli download mistralai/Mistral-7B-Instruct-v0.3 --local-dir D:\\LLM_Models\\Mistral-7B-Instruct-v0.3"""

    add_code_block(doc, download_commands)

    doc.add_paragraph()

    # 15.8 Limitations
    add_heading(doc, "15.8 Limitations du mode offline", level=2)

    add_bullet_list(doc, [
        "Ingestion de nouveaux documents non disponible (requiert les embeddings API)",
        "Confluence non disponible (requiert connexion reseau)",
        "Qualite des reponses legerement inferieure au mode online",
        "Utilise la VRAM en permanence (~12 GB)",
    ])

    # Sauvegarder le document
    doc.save('GUIDE_RAG.docx')
    print("Document GUIDE_RAG.docx mis a jour avec la section Mode OFFLINE (v2.1)")

if __name__ == "__main__":
    main()
